#!/usr/bin/env python3
"""
Word MCP Server - FastMCP version.python-docx-based tools.
"""

import os
import shutil
from typing import Any, Dict, List, Optional, Tuple
from fastmcp import FastMCP
from dotenv import load_dotenv
import logging

# Load environment variables from .env file
load_dotenv()
from docx import Document  # pip install python-docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Optional PDF export
try:
    from docx2pdf import convert as docx2pdf_convert  # pip install docx2pdf
    HAS_DOCX2PDF = True
except Exception:
    HAS_DOCX2PDF = False

logging.basicConfig(level=os.getenv("LOG_LEVEL", "INFO"))
DRY = os.getenv("DRY_RUN", "0") == "true"
BASE_DIR = os.getenv("WORD_DOCS_BASE_DIR", "")

mcp = FastMCP("Word MCP (native)")

def _p(path: str) -> str:
    return os.path.join(BASE_DIR, path) if BASE_DIR and not os.path.isabs(path) else path

def _dry(name: str, **kwargs):
    logging.info("DRY: %s %s", name, kwargs)
    return {"dry_run": True, "tool": f"word_{name}", "args": kwargs}

# ---------- Document Management ----------

@mcp.tool()
def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> Dict[str, Any]:
    """Create a new Word document with optional metadata."""
    if DRY:
        return _dry("create_document", filename=filename, title=title, author=author)
    doc = Document()
    if title: doc.core_properties.title = title
    if author: doc.core_properties.author = author
    path = _p(filename)
    doc.save(path)
    return {"status": "success", "filename": path}

@mcp.tool()
def get_document_info(filename: str) -> Dict[str, Any]:
    """Get document core properties and simple stats."""
    if DRY:
        return _dry("get_document_info", filename=filename)
    doc = Document(_p(filename))
    props = doc.core_properties
    stats = {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}
    return {
        "title": props.title, "author": props.author, "subject": props.subject,
        "created": str(props.created), "last_modified_by": props.last_modified_by,
        "stats": stats
    }

@mcp.tool()
def list_available_documents(directory: str = ".") -> Dict[str, Any]:
    """List .docx files in a directory."""
    if DRY:
        return _dry("list_available_documents", directory=directory)
    dirp = _p(directory)
    files = [f for f in os.listdir(dirp) if f.lower().endswith(".docx")]
    return {"directory": dirp, "files": files}

@mcp.tool()
def copy_document(source_filename: str, destination_filename: Optional[str] = None) -> Dict[str, Any]:
    """Copy a document to a new file."""
    if DRY:
        return _dry("copy_document", source_filename=source_filename, destination_filename=destination_filename)
    src = _p(source_filename)
    dst = _p(destination_filename or f"copy_{os.path.basename(source_filename)}")
    shutil.copyfile(src, dst)
    return {"status": "success", "source": src, "destination": dst}

@mcp.tool()
def merge_documents(filenames: List[str], output_filename: str) -> Dict[str, Any]:
    """Merge multiple documents into a single document (simple append)."""
    if DRY:
        return _dry("merge_documents", filenames=filenames, output_filename=output_filename)
    base = Document(_p(filenames))
    for name in filenames[1:]:
        src = Document(_p(name))
        for element in src.element.body:
            base.element.body.append(element)
    out = _p(output_filename)
    base.save(out)
    return {"status": "success", "output": out}

@mcp.tool()
def convert_to_pdf(filename: str, output_filename: Optional[str] = None) -> Dict[str, Any]:
    """Convert DOCX to PDF using docx2pdf if available."""
    if DRY:
        return _dry("convert_to_pdf", filename=filename, output_filename=output_filename)
    if not HAS_DOCX2PDF:
        return {"status": "error", "message": "docx2pdf not available (requires Word on Windows/macOS)"}
    src = _p(filename)
    out = _p(output_filename or os.path.splitext(filename) + ".pdf")
    docx2pdf_convert(src, out)
    return {"status": "success", "output": out}

# ---------- Content Addition ----------

@mcp.tool()
def add_heading(filename: str, text: str, level: int = 1) -> Dict[str, Any]:
    if DRY:
        return _dry("add_heading", filename=filename, text=text, level=level)
    doc = Document(_p(filename))
    doc.add_heading(text, level=level)
    doc.save(_p(filename))
    return {"status": "success"}

@mcp.tool()
def add_paragraph(filename: str, text: str, style: Optional[str] = None, align: Optional[str] = None) -> Dict[str, Any]:
    if DRY:
        return _dry("add_paragraph", filename=filename, text=text, style=style, align=align)
    doc = Document(_p(filename))
    p = doc.add_paragraph(text)
    if style: p.style = style
    if align:
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, align.upper(), None)
    doc.save(_p(filename))
    return {"status": "success"}

@mcp.tool()
def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> Dict[str, Any]:
    if DRY:
        return _dry("add_table", filename=filename, rows=rows, cols=cols, data=data)
    doc = Document(_p(filename))
    table = doc.add_table(rows=rows, cols=cols)
    if data:
        for r in range(min(rows, len(data))):
            for c in range(min(cols, len(data[r]))):
                table.cell(r, c).text = str(data[r][c])
    doc.save(_p(filename))
    return {"status": "success"}

@mcp.tool()
def add_picture(filename: str, image_path: str, width_inches: Optional[float] = None) -> Dict[str, Any]:
    if DRY:
        return _dry("add_picture", filename=filename, image_path=image_path, width_inches=width_inches)
    doc = Document(_p(filename))
    width = Inches(width_inches) if width_inches else None
    doc.add_picture(_p(image_path), width=width)
    doc.save(_p(filename))
    return {"status": "success"}

@mcp.tool()
def add_page_break(filename: str) -> Dict[str, Any]:
    if DRY:
        return _dry("add_page_break", filename=filename)
    doc = Document(_p(filename))
    doc.add_page_break()
    doc.save(_p(filename))
    return {"status": "success"}

# ---------- Content Extraction ----------

@mcp.tool()
def get_document_text(filename: str) -> Dict[str, Any]:
    if DRY:
        return _dry("get_document_text", filename=filename)
    doc = Document(_p(filename))
    text = "\n".join(p.text for p in doc.paragraphs)
    return {"text": text}

@mcp.tool()
def get_paragraph_text_from_document(filename: str, paragraph_index: int) -> Dict[str, Any]:
    if DRY:
        return _dry("get_paragraph_text_from_document", filename=filename, paragraph_index=paragraph_index)
    doc = Document(_p(filename))
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        return {"status": "error", "message": "paragraph_index out of range"}
    return {"text": doc.paragraphs[paragraph_index].text}

@mcp.tool()
def find_text_in_document(filename: str, text_to_find: str, match_case: bool = True, whole_word: bool = False) -> Dict[str, Any]:
    if DRY:
        return _dry("find_text_in_document", filename=filename, text_to_find=text_to_find, match_case=match_case, whole_word=whole_word)
    doc = Document(_p(filename))
    matches: List[Tuple[int, int]] = []
    needle = text_to_find if match_case else text_to_find.lower()
    for i, p in enumerate(doc.paragraphs):
        hay = p.text if match_case else p.text.lower()
        idx = 0
        while True:
            pos = hay.find(needle, idx)
            if pos == -1:
                break
            if whole_word:
                before = hay[pos - 1] if pos > 0 else " "
                after = hay[pos + len(needle)] if pos + len(needle) < len(hay) else " "
                if before.isalnum() or after.isalnum():
                    idx = pos + 1
                    continue
            matches.append((i, pos))
            idx = pos + len(needle)
    return {"matches": [{"paragraph_index": pi, "char_index": ci} for pi, ci in matches]}

# ---------- Rich Text Formatting ----------

@mcp.tool()
def format_text(filename: str, paragraph_index: int, start_pos: int, end_pos: int,
                bold: Optional[bool] = None, italic: Optional[bool] = None, underline: Optional[bool] = None,
                color: Optional[str] = None, font_size: Optional[float] = None, font_name: Optional[str] = None) -> Dict[str, Any]:
    if DRY:
        return _dry("format_text", filename=filename, paragraph_index=paragraph_index, start_pos=start_pos, end_pos=end_pos)
    doc = Document(_p(filename))
    p = doc.paragraphs[paragraph_index]
    # Rebuild runs around target span
    text = p.text
    p.clear()  # remove existing runs
    before = p.add_run(text[:start_pos])
    target = p.add_run(text[start_pos:end_pos])
    after = p.add_run(text[end_pos:])
    if bold is not None: target.bold = bold
    if italic is not None: target.italic = italic
    if underline is not None: target.underline = underline
    if color: target.font.color.rgb = bytes.fromhex(color)
    if font_size: target.font.size = Pt(font_size)
    if font_name: target.font.name = font_name
    doc.save(_p(filename))
    return {"status": "success"}

@mcp.tool()
def search_and_replace(filename: str, find_text: str, replace_text: str) -> Dict[str, Any]:
    if DRY:
        return _dry("search_and_replace", filename=filename, find_text=find_text, replace_text=replace_text)
    doc = Document(_p(filename))
    for p in doc.paragraphs:
        if find_text in p.text:
            p.text = p.text.replace(find_text, replace_text)
    doc.save(_p(filename))
    return {"status": "success"}

@mcp.tool()
def delete_paragraph(filename: str, paragraph_index: int) -> Dict[str, Any]:
    if DRY:
        return _dry("delete_paragraph", filename=filename, paragraph_index=paragraph_index)
    doc = Document(_p(filename))
    p = doc.paragraphs[paragraph_index]
    p._element.getparent().remove(p._element)
    doc.save(_p(filename))
    return {"status": "success"}

@mcp.tool()
def create_custom_style(filename: str, style_name: str, bold: Optional[bool] = None, italic: Optional[bool] = None,
                        font_size: Optional[float] = None, font_name: Optional[str] = None, color: Optional[str] = None,
                        base_style: Optional[str] = None) -> Dict[str, Any]:
    if DRY:
        return _dry("create_custom_style", filename=filename, style_name=style_name)
    doc = Document(_p(filename))
    styles = doc.styles
    style = styles.add_style(style_name, 1)  # 1=character style
    if base_style and base_style in styles:
        style.base_style = styles[base_style]
    font = style.font
    if bold is not None: font.bold = bold
    if italic is not None: font.italic = italic
    if font_size: font.size = Pt(font_size)
    if font_name: font.name = font_name
    if color: font.color.rgb = bytes.fromhex(color)
    doc.save(_p(filename))
    return {"status": "success"}

# ---------- Table Formatting ----------

def _cell_shade(cell, hex_color: str):
    # Add w:shd element to cell properties (fill color) [4]
    shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), hex_color))
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(shading)

@mcp.tool()
def format_table(filename: str, table_index: int, has_header_row: Optional[bool] = None,
                 border_style: Optional[str] = None, shading: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
    if DRY:
        return _dry("format_table", filename=filename, table_index=table_index, has_header_row=has_header_row, border_style=border_style, shading=shading)
    doc = Document(_p(filename))
    tbl = doc.tables[table_index]
    if has_header_row is not None:
        tbl.rows.cells.paragraphs.style = "Heading 2" if has_header_row else tbl.rows.cells.paragraphs.style
    if shading and "row" in shading and "color" in shading:
        r = int(shading["row"])
        col = shading.get("col")
        if col is None:
            for c in tbl.rows[r].cells:
                _cell_shade(c, shading["color"])
        else:
            _cell_shade(tbl.cell(r, int(col)), shading["color"])
    doc.save(_p(filename))
    return {"status": "success"}

@mcp.tool()
def set_table_cell_shading(filename: str, table_index: int, row_index: int, col_index: int, fill_color: str) -> Dict[str, Any]:
    if DRY:
        return _dry("set_table_cell_shading", filename=filename, table_index=table_index, row_index=row_index, col_index=col_index, fill_color=fill_color)
    doc = Document(_p(filename))
    tbl = doc.tables[table_index]
    _cell_shade(tbl.cell(row_index, col_index), fill_color)
    doc.save(_p(filename))
    return {"status": "success"}

@mcp.tool()
def merge_table_cells(filename: str, table_index: int, start_row: int, start_col: int, end_row: int, end_col: int) -> Dict[str, Any]:
    if DRY:
        return _dry("merge_table_cells", filename=filename, table_index=table_index, start_row=start_row, start_col=start_col, end_row=end_row, end_col=end_col)
    doc = Document(_p(filename))
    tbl = doc.tables[table_index]
    a = tbl.cell(start_row, start_col)
    b = tbl.cell(end_row, end_col)
    a.merge(b)  # supports horiz/vert merge [3]
    doc.save(_p(filename))
    return {"status": "success"}

# ---------- Comment Extraction (placeholder: python-docx has limited comments support) ----------

@mcp.tool()
def get_all_comments(filename: str) -> Dict[str, Any]:
    """Extract comments metadata if present; python-docx has limited native support."""
    if DRY:
        return _dry("get_all_comments", filename=filename)
    # python-docx does not expose comments API; would require raw XML parsing.
    return {"status": "error", "message": "Comment extraction requires custom XML parsing not included here."}

if __name__ == "__main__":
    mcp.run()
